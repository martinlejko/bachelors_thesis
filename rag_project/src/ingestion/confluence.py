"""
Confluence data ingestion module.

This module handles fetching DOCX files from Confluence and converting them to the
Document model format using the Confluence API.
"""

import os
import logging
import requests
import tempfile
from datetime import datetime
from typing import List, Dict, Any, Optional
import docx

from src.common.models import Document, DocumentSource
from src.common.config import CONFLUENCE_URL, CONFLUENCE_USERNAME, CONFLUENCE_API_KEY, CACHE_DIR
from src.common.utils import create_hash, save_to_cache, load_from_cache
from src.ingestion.base import DataIngestionSource

logger = logging.getLogger(__name__)


class ConfluenceSource(DataIngestionSource):
    """Data ingestion source for Confluence."""

    def __init__(self, space_key: str, limit: int = 100):
        """
        Initialize the Confluence source.

        Args:
            space_key: The key of the Confluence space
            limit: Maximum number of pages to fetch
        """
        self.space_key = space_key
        self.limit = limit
        self.auth = (CONFLUENCE_USERNAME, CONFLUENCE_API_KEY)
        self.cache_dir = CACHE_DIR / "confluence"
        os.makedirs(self.cache_dir, exist_ok=True)

        if not CONFLUENCE_URL or not CONFLUENCE_USERNAME or not CONFLUENCE_API_KEY:
            raise ValueError("Confluence API credentials not configured.")

    def get_source_type(self) -> DocumentSource:
        """Get the source type."""
        return DocumentSource.CONFLUENCE

    def get_source_info(self) -> Dict[str, Any]:
        """Get information about the source."""
        return {
            "space_key": self.space_key,
            "limit": self.limit,
            "url": CONFLUENCE_URL,
            "timestamp": datetime.now().isoformat(),
        }

    def has_changed(self) -> bool:
        """Check if the Confluence space has changed."""
        try:
            current_pages = self._get_page_list()

            cache_file = f"confluence_{self.space_key}_pages.json"
            cached_pages = load_from_cache(self.cache_dir, cache_file)

            if cached_pages is None:
                logger.info("No cached page list found, considering as changed")
                return True

            current_hash = create_hash(current_pages)
            cached_hash = create_hash(cached_pages)

            changed = current_hash != cached_hash
            if changed:
                logger.info("Confluence pages have changed")
            else:
                logger.info("Confluence pages have not changed")

            return changed

        except Exception as e:
            logger.error(f"Error checking if Confluence space has changed: {e}")
            return True

    def load_data(self) -> List[Document]:
        """Load data from Confluence."""
        try:
            pages = self._get_page_list()

            cache_file = f"confluence_{self.space_key}_pages.json"
            save_to_cache(pages, self.cache_dir, cache_file)

            documents = []

            for page in pages:
                try:
                    page_id = page["id"]
                    page_title = page["title"]

                    docx_content = self._download_page_as_docx(page_id)

                    if docx_content:
                        text = self._extract_text_from_docx(docx_content)

                        doc_id = f"confluence_{page_id}"
                        metadata = {
                            "title": page_title,
                            "id": page_id,
                            "space_key": self.space_key,
                            "url": f"{CONFLUENCE_URL}/pages/viewpage.action?pageId={page_id}",
                        }

                        document = Document(
                            id=doc_id, content=text, metadata=metadata, source=DocumentSource.CONFLUENCE
                        )

                        documents.append(document)
                except Exception as e:
                    logger.error(f"Error processing Confluence page {page.get('id')}: {e}")

            logger.info(f"Loaded {len(documents)} documents from Confluence")
            return documents

        except Exception as e:
            logger.error(f"Error loading data from Confluence: {e}")
            return []

    def _get_page_list(self) -> List[Dict[str, Any]]:
        """
        Get the list of pages from the Confluence space.

        Returns:
            List[Dict[str, Any]]: List of pages
        """
        url = f"{CONFLUENCE_URL}/rest/api/content"
        params = {"spaceKey": self.space_key, "limit": self.limit, "expand": "version"}

        response = requests.get(url, params=params, auth=self.auth)
        response.raise_for_status()

        data = response.json()
        return data.get("results", [])

    def _download_page_as_docx(self, page_id: str) -> Optional[bytes]:
        """
        Download a page as DOCX.

        Args:
            page_id: ID of the page

        Returns:
            Optional[bytes]: DOCX content or None if failed
        """
        url = f"{CONFLUENCE_URL}/exportword?pageId={page_id}"

        response = requests.get(url, auth=self.auth)
        if response.status_code == 200:
            return response.content
        else:
            logger.error(f"Failed to download page {page_id} as DOCX: {response.status_code}")
            return None

    def _extract_text_from_docx(self, docx_content: bytes) -> str:
        """
        Extract text from DOCX content.

        Args:
            docx_content: DOCX content as bytes

        Returns:
            str: Extracted text
        """
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
            temp_file.write(docx_content)
            temp_file_path = temp_file.name

        try:
            doc = docx.Document(temp_file_path)
            text = "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)
            return text
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return ""
        finally:
            if os.path.exists(temp_file_path):
                os.unlink(temp_file_path)
